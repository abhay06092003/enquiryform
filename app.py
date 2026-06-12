import io
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from flask import (
    Flask,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)

from storage import get_submission, init_db, insert_submission, list_submissions, storage_backend_name

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "quick-platter-enquiry-secret-2026")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "abhay123")

init_db()


def is_admin():
    return session.get("admin_authenticated") is True


def require_admin():
    return is_admin()


def format_responses(responses):
    if isinstance(responses, str):
        return json.loads(responses)
    return responses


def group_responses_by_section(responses):
    sections = []
    current = None
    for item in responses:
        if item.get("type") == "section":
            current = {"title": item["label"], "questions": []}
            sections.append(current)
        elif current is not None:
            current["questions"].append(item)
    return sections


def build_docx(name, responses, submitted_at):
    doc = Document()

    title = doc.add_heading("Quick Platter — Requirement Gathering Questionnaire", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    meta = doc.add_paragraph()
    meta.add_run("Respondent: ").bold = True
    meta.add_run(name)
    meta.add_run("\nSubmitted: ").bold = True
    meta.add_run(submitted_at)

    doc.add_paragraph()

    for item in responses:
        if item.get("type") == "section":
            doc.add_heading(item["label"], level=1)
            continue

        p = doc.add_paragraph()
        q_run = p.add_run(f"{item['id']}. {item['question']}")
        q_run.bold = True
        q_run.font.size = Pt(11)

        answer = item.get("answer", "")
        if isinstance(answer, list):
            answer_text = ", ".join(answer) if answer else "—"
        elif answer:
            answer_text = str(answer)
        else:
            answer_text = "—"

        ans_p = doc.add_paragraph(answer_text)
        ans_p.paragraph_format.left_indent = Pt(18)
        ans_p.runs[0].font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/submit", methods=["POST"])
def submit():
    data = request.get_json()
    if not data or not data.get("name", "").strip():
        return jsonify({"success": False, "error": "Name is required."}), 400

    name = data["name"].strip()
    responses = data.get("responses", [])
    submitted_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    insert_submission(name, responses, submitted_at)

    return jsonify({"success": True, "message": "Your responses have been submitted successfully!"})


@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    error = None
    if request.method == "POST":
        password = request.form.get("password", "")
        if password == ADMIN_PASSWORD:
            session["admin_authenticated"] = True
            return redirect(url_for("admin_dashboard"))
        error = "Incorrect password. Please try again."
    return render_template("admin_login.html", error=error)


@app.route("/admin/logout")
def admin_logout():
    session.pop("admin_authenticated", None)
    return redirect(url_for("admin_login"))


@app.route("/admin")
def admin_dashboard():
    if not require_admin():
        return redirect(url_for("admin_login"))

    submissions = list_submissions()
    return render_template(
        "admin.html",
        submissions=submissions,
        storage_backend=storage_backend_name(),
    )


@app.route("/admin/submission/<int:submission_id>")
def admin_view_submission(submission_id):
    if not require_admin():
        return redirect(url_for("admin_login"))

    row = get_submission(submission_id)
    if not row:
        return "Submission not found.", 404

    submission = dict(row)
    submission["sections"] = group_responses_by_section(
        format_responses(submission["responses"])
    )
    return render_template("admin_view.html", submission=submission)


@app.route("/admin/download/<int:submission_id>")
def admin_download(submission_id):
    if not require_admin():
        return redirect(url_for("admin_login"))

    row = get_submission(submission_id)
    if not row:
        return "Submission not found.", 404

    responses = format_responses(row["responses"])
    buffer = build_docx(row["name"], responses, row["submitted_at"])

    safe_name = "".join(c if c.isalnum() or c in " -_" else "_" for c in row["name"])
    filename = f"Quick_Platter_{safe_name}_{submission_id}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
